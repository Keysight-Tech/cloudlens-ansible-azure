#!/usr/bin/env python3
"""
CloudLens Stack Deployment Runbook generator.

Builds CloudLens_Stack_Deployment_Runbook.docx using the same visual
language as generate_runbook.py:
  - Keysight blue (#0078D4) section headers
  - Azure-blue header rows on tables
  - Courier New code blocks
  - Calibri body, professional sans-serif
  - Cover page + static (LibreOffice-safe) TOC + page numbers

Audience: customer DevOps, CTO/procurement, training. Covers the three
deployment paths (Bash one-liner / Terraform stack module / Azure
Portal click-through) plus prerequisites, verification, troubleshooting
and cleanup.

Run:
    python3 docs/generate_stack_runbook.py
Output:
    docs/CloudLens_Stack_Deployment_Runbook.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# -----------------------------------------------------------------------------
# Brand palette (matches generate_runbook.py exactly)
# -----------------------------------------------------------------------------
AZURE_BLUE = RGBColor(0x00, 0x78, 0xD4)
AZURE_DARK = RGBColor(0x00, 0x5A, 0x9E)
KEYSIGHT_GOLD = RGBColor(0xD4, 0xAF, 0x37)
KEYSIGHT_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEXT_DARK = RGBColor(0x1B, 0x2A, 0x4A)
TEXT_MUTED = RGBColor(0x55, 0x65, 0x75)
CALLOUT_BG = "DDEBF7"
CODE_BG = "F2F2F2"
TABLE_HEADER_BG = "0078D4"
TABLE_ALT_ROW_BG = "F4F8FB"
AZURE_BLUE_HEX = "0078D4"

# -----------------------------------------------------------------------------
# Paths
# -----------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
OUTPUT_DOCX = SCRIPT_DIR / "CloudLens_Stack_Deployment_Runbook.docx"


# -----------------------------------------------------------------------------
# XML helpers
# -----------------------------------------------------------------------------
def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_MUTED


# Static, pre-evaluated TOC: kept in sync with build_* calls in main().
TOC_ENTRIES: list[tuple[int, str]] = [
    (1, "1. Executive Summary"),
    (1, "2. The Three Deployment Paths"),
    (1, "3. Prerequisites"),
    (1, "4. Path A: Bash One-Liner"),
    (2, "4.1  What the script does (phase by phase)"),
    (2, "4.2  Flags and overrides"),
    (1, "5. Path B: Terraform Stack Module"),
    (2, "5.1  Workflow"),
    (2, "5.2  What gets created"),
    (1, "6. Path C: Azure Portal"),
    (1, "7. Verification Checklist"),
    (1, "8. Troubleshooting Reference"),
    (1, "9. Cleanup / Decommission"),
    (1, "10. Appendix: File Paths and Quick Links"),
]


def _add_toc_entry(doc: Document, level: int, label: str) -> None:
    """Add one TOC line: indented by level, leader-dot tab stop on the right."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    if level >= 2:
        pf.left_indent = Inches(0.35)

    tab_stops = pf.tab_stops
    try:
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    except Exception:
        tab_stops.add_tab_stop(Inches(6.3))

    label_run = p.add_run(label)
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(12) if level == 1 else Pt(10.5)
    label_run.bold = (level == 1)
    label_run.font.color.rgb = AZURE_DARK if level == 1 else TEXT_DARK

    # Trailing tab pushes leader dots to the right margin.
    p.add_run("\t")


# -----------------------------------------------------------------------------
# Style configuration
# -----------------------------------------------------------------------------
def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = AZURE_BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = AZURE_DARK
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = KEYSIGHT_NAVY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True


# -----------------------------------------------------------------------------
# Reusable element builders
# -----------------------------------------------------------------------------
def add_paragraph(doc: Document, text: str, *, bold: bool = False,
                  italic: bool = False, color: RGBColor | None = None,
                  size: int = 11, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_code_block(doc: Document, code: str) -> None:
    """Monospace Courier New block with light shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)

    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_BG)
    p_pr.append(shd)

    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = KEYSIGHT_NAVY


def add_callout(doc: Document, label: str, body: str) -> None:
    """Single-cell shaded callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CALLOUT_BG)
    _set_cell_borders(cell, color=AZURE_BLUE_HEX, size="6")

    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(2)
    run = label_p.add_run(label)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = AZURE_DARK

    body_p = cell.add_paragraph()
    run2 = body_p.add_run(body)
    run2.font.name = "Calibri"
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = TEXT_DARK

    doc.add_paragraph()


def add_checkbox_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        box = p.add_run("☐  ")
        box.font.name = "Segoe UI Symbol"
        box.font.size = Pt(13)
        box.font.color.rgb = AZURE_DARK
        text = p.add_run(item)
        text.font.name = "Calibri"
        text.font.size = Pt(11)
        text.font.color.rgb = TEXT_DARK


def add_numbered(doc: Document, steps: list[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                     col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade_cell(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            cell = row_cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                _shade_cell(cell, TABLE_ALT_ROW_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_DARK

    if col_widths:
        for row in table.rows:
            for c_idx, width_in in enumerate(col_widths):
                row.cells[c_idx].width = Inches(width_in)

    doc.add_paragraph()


# -----------------------------------------------------------------------------
# Header / Footer
# -----------------------------------------------------------------------------
def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.different_first_page_header_footer = True

    header = section.header
    h_p = header.paragraphs[0]
    h_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = h_p.add_run("CloudLens Stack Deployment Runbook")
    h_run.font.name = "Calibri"
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = AZURE_DARK
    h_run.italic = True

    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = f_p.add_run("Keysight Technologies   |   ")
    left.font.name = "Calibri"
    left.font.size = Pt(9)
    left.font.color.rgb = TEXT_MUTED
    _add_page_number(f_p)


# -----------------------------------------------------------------------------
# Cover page
# -----------------------------------------------------------------------------
def build_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run("KEYSIGHT")
    logo_run.font.name = "Calibri"
    logo_run.font.size = Pt(20)
    logo_run.bold = True
    logo_run.font.color.rgb = KEYSIGHT_GOLD

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline_p.add_run("TECHNOLOGIES")
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = KEYSIGHT_NAVY
    tag_run.bold = True

    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CloudLens Stack")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(36)
    title_run.bold = True
    title_run.font.color.rgb = AZURE_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Deployment Runbook")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(26)
    sub_run.font.color.rgb = AZURE_DARK

    doc.add_paragraph()
    doc.add_paragraph()

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag_p.add_run(
        "CLMS, vPB and sensors. End to end. One command."
    )
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(14)
    tag_run.italic = True
    tag_run.font.color.rgb = TEXT_MUTED

    for _ in range(8):
        doc.add_paragraph()

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    for row in meta_table.rows:
        for cell in row.cells:
            cell.width = Inches(2.0)
    labels = [("Version", "v1.0"), ("Date", "June 2026")]
    for r_idx, (label, value) in enumerate(labels):
        l_cell = meta_table.rows[r_idx].cells[0]
        v_cell = meta_table.rows[r_idx].cells[1]
        _shade_cell(l_cell, "0078D4")
        _shade_cell(v_cell, "F4F8FB")
        lp = l_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        lr.font.name = "Calibri"
        lr.font.size = Pt(11)
        vp = v_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(value)
        vr.bold = True
        vr.font.color.rgb = KEYSIGHT_NAVY
        vr.font.name = "Calibri"
        vr.font.size = Pt(11)

    for _ in range(2):
        doc.add_paragraph()

    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot_p.add_run("Keysight Technologies   |   Network Visibility Solutions")
    foot_run.font.name = "Calibri"
    foot_run.font.size = Pt(10)
    foot_run.italic = True
    foot_run.font.color.rgb = TEXT_MUTED

    doc.add_page_break()


# -----------------------------------------------------------------------------
# Sections
# -----------------------------------------------------------------------------
def build_toc(doc: Document) -> None:
    """Static TOC. We avoid Word's TOC field so the LibreOffice PDF converter
    renders the entries directly instead of leaving a placeholder.
    """
    doc.add_heading("Table of Contents", level=1)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    for level, label in TOC_ENTRIES:
        _add_toc_entry(doc, level, label)
    doc.add_page_break()


def build_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    add_paragraph(
        doc,
        "This runbook shows how to deploy the full CloudLens stack on Azure in "
        "one shot: CloudLens Manager (CLMS), optional Virtual Packet Broker "
        "(vPB), and the OS sensors that connect tagged VMs back to CLMS. "
        "Three paths are documented; pick the one that matches how your team "
        "works. All three produce the same Azure resources.",
    )

    doc.add_heading("Audience", level=2)
    add_bullets(doc, [
        "Customer DevOps and platform engineers who own the Azure subscription.",
        "Keysight Sales Engineers running customer POCs or demos.",
        "Customer CTO / procurement / training teams who need a printable reference.",
    ])

    doc.add_heading("What gets deployed", level=2)
    add_styled_table(
        doc,
        headers=["Component", "Azure resources", "Required"],
        rows=[
            ["CloudLens Manager (CLMS)",
             "1 VM (D4s_v5), 1 public IP, NSG, NIC",
             "Yes"],
            ["Virtual Packet Broker (vPB)",
             "1 VM (D8s_v3), 1 public IP for mgmt, NSG, 3 NICs",
             "Optional"],
            ["Sensors",
             "Agent containers (Linux) and native service (Windows) on tagged VMs",
             "Optional"],
            ["Shared networking",
             "VNet 10.50.0.0/16, 1-4 subnets depending on options",
             "Yes"],
        ],
        col_widths=[2.0, 3.4, 1.2],
    )

    doc.add_heading("Time and cost", level=2)
    add_styled_table(
        doc,
        headers=["Phase", "Wall-clock time"],
        rows=[
            ["CLMS VM provisioned and initialised", "15 to 20 min"],
            ["vPB VM provisioned (optional)", "5 to 10 min"],
            ["Sensors deployed on first 10 VMs (optional)", "5 to 8 min"],
            ["Total typical (CLMS + vPB + sensors)", "30 to 40 min"],
        ],
        col_widths=[3.6, 2.4],
    )
    add_paragraph(
        doc,
        "Indicative Azure cost (East US 2, June 2026 list prices): CLMS "
        "D4s_v5 about USD 140/month, vPB D8s_v3 about USD 280/month, plus "
        "small storage and public-IP fees. Stopping the VMs deallocates compute "
        "charges. Marketplace software fees are billed separately by Keysight.",
        italic=True, size=10, color=TEXT_MUTED,
    )
    doc.add_page_break()


def build_paths_overview(doc: Document) -> None:
    doc.add_heading("2. The Three Deployment Paths", level=1)
    add_paragraph(
        doc,
        "All three paths accept Azure Marketplace terms automatically, deploy "
        "the same VMs, and chain CLMS first, then vPB. Pick the one that "
        "matches your operating model. You can switch between them later "
        "without losing state, because each path produces standard Azure "
        "resources in a resource group.",
    )
    add_styled_table(
        doc,
        headers=["Path", "Best for", "Tools needed", "Time to first paste"],
        rows=[
            ["A. Bash one-liner",
             "First-time trial, customer demo, Cloud Shell users",
             "Browser + Azure CLI (Cloud Shell has both)",
             "10 seconds"],
            ["B. Terraform stack module",
             "IaC pipelines, repeatable customer envs, GitOps",
             "Terraform 1.5+ and Azure CLI logged in",
             "2 minutes"],
            ["C. Azure Portal",
             "Click-through customers, governance-controlled subscriptions",
             "Browser only",
             "1 minute"],
        ],
        col_widths=[1.6, 2.6, 1.9, 1.3],
    )
    add_callout(
        doc, "Why three?",
        "Customers do not all work the same way. A platform team wants Terraform. "
        "A network engineer trialling the kit wants a paste. A CTO walking through "
        "with procurement wants buttons. Three paths, identical outcome.",
    )
    doc.add_page_break()


def build_prerequisites(doc: Document) -> None:
    doc.add_heading("3. Prerequisites", level=1)
    add_paragraph(
        doc,
        "Confirm each item before kicking off the deployment. The script will "
        "stop early with a clear message if any of the first three are missing, "
        "so you do not waste time mid-deploy.",
    )
    add_checkbox_list(doc, [
        "Azure subscription with Contributor role on the target resource group "
        "(or rights to create one).",
        "Azure CLI 2.50 or later installed and logged in. In Cloud Shell this "
        "is preinstalled and pre-authenticated.",
        "Marketplace terms can be accepted automatically by the script; you do "
        "not need to click through anything in the Portal first.",
        "Quota for the DSv5 family (CLMS, 4 vCPU) and DSv3 family (vPB, 8 vCPU) "
        "in your chosen region. The script does a soft quota probe and warns "
        "without blocking.",
        "Open egress on the VM subnets to reach the CLMS public IP on TCP/443 "
        "(only needed once sensors are deployed).",
    ])
    add_callout(
        doc, "Region note",
        "Default region is eastus2. CLMS and vPB Marketplace images are also "
        "published in westus2, northeurope, westeurope, southeastasia. Run "
        "'az vm image list --offer keysight-cloudlens-manager-preview "
        "--publisher keysight-technologies-cloudlens --all' to confirm before "
        "trying a less-common region.",
    )
    doc.add_page_break()


def build_path_a(doc: Document) -> None:
    doc.add_heading("4. Path A: Bash One-Liner", level=1)
    add_paragraph(
        doc,
        "The fastest way to a working stack. One paste in Azure Cloud Shell or "
        "any local terminal with the Azure CLI; the script handles the rest.",
    )
    add_code_block(
        doc,
        "curl -sSL "
        "https://raw.githubusercontent.com/Keysight-Tech/"
        "cloudlens-ansible-azure/main/deploy/deploy-stack.sh | bash",
    )
    add_paragraph(
        doc,
        "Prefer to inspect before running? Download then execute:",
    )
    add_code_block(
        doc,
        "curl -sSL -o deploy-stack.sh \\\n"
        "  https://raw.githubusercontent.com/Keysight-Tech/"
        "cloudlens-ansible-azure/main/deploy/deploy-stack.sh\n"
        "less deploy-stack.sh\n"
        "bash deploy-stack.sh",
    )

    doc.add_heading("4.1  What the script does (phase by phase)", level=2)
    add_paragraph(
        doc,
        "The script runs 10 phases. Each phase prints a banner with its name, "
        "so if anything fails you know exactly where you are.",
    )
    add_numbered(doc, [
        "Banner and environment detection: prints the version header and "
        "detects Cloud Shell vs local terminal so prompts adapt accordingly.",
        "Pre-flight checks: confirms az CLI is installed, you are logged in, "
        "and probes vCPU quota for the DSv5 and DSv3 families.",
        "Customer input: prompts for resource group name (default cloudlens-rg), "
        "Azure region (default eastus2), and a strong admin password. All "
        "defaults can be overridden by flags (see 4.2).",
        "Marketplace terms acceptance: runs 'az vm image terms accept' for the "
        "CLMS and (if selected) vPB plans so the deployments do not stall on "
        "marketplace consent.",
        "Resource group creation: creates the RG only if it does not exist. "
        "Re-running the script against an existing RG is safe.",
        "CLMS deployment: deploys the clms-marketplace.json ARM template. "
        "Waits for the VM provisioning state to be Succeeded and captures the "
        "public IP.",
        "Wait for CLMS init: polls the CLMS HTTPS endpoint until the UI is "
        "reachable (typically 15 minutes). Prints a one-line progress bar.",
        "vPB deployment (optional): the prompt defaults to Yes; type N to skip. "
        "When enabled, deploys vpb-marketplace.json and captures the management "
        "public IP.",
        "Manual project key step: the script pauses and tells you to open the "
        "CLMS UI, create a project, and paste the project key back into the "
        "terminal. This is the one step that cannot be automated end to end "
        "today (CLMS does not yet expose a project-create REST endpoint).",
        "Sensor chain (optional): hands off to quickstart.sh, which deploys "
        "sensors to every Azure VM tagged cloudlens=yes. Prints a final summary "
        "and writes cloudlens-deploy-summary.txt with all IPs and credentials.",
    ])
    add_callout(
        doc, "Output you can keep",
        "Every run writes 'cloudlens-deploy-stack.log' (raw stdout/stderr) and "
        "'cloudlens-deploy-summary.txt' (RG name, VM names, IPs, default creds). "
        "Save both to your customer ticket or runbook archive.",
    )

    doc.add_heading("4.2  Flags and overrides", level=2)
    add_styled_table(
        doc,
        headers=["Flag", "Effect"],
        rows=[
            ["--dry-run",
             "Walk through every prompt and print the az commands without "
             "touching Azure. Use this to preview the run."],
            ["--resource-group NAME",
             "Override the default RG name (cloudlens-rg). Useful when "
             "deploying into an existing customer RG."],
            ["--location REGION",
             "Override the default region (eastus2). Pass any region that "
             "carries the Marketplace images."],
            ["--no-vpb",
             "Skip the vPB phase entirely. CLMS-only deploys complete in "
             "about 20 minutes."],
            ["--no-sensors",
             "Stop after the CLMS (and vPB) phase. Run quickstart.sh later "
             "yourself if you want sensors then."],
            ["-h | --help",
             "Print the help banner with all flags and a phase summary."],
        ],
        col_widths=[2.0, 4.6],
    )
    doc.add_page_break()


def build_path_b(doc: Document) -> None:
    doc.add_heading("5. Path B: Terraform Stack Module", level=1)
    add_paragraph(
        doc,
        "The stack module wraps the existing clms and vpb modules behind one "
        "tfvars file. One 'terraform apply' provisions the resource group, a "
        "shared VNet with the right subnets, both VMs, NSGs, NICs and public "
        "IPs. The sensors are still deployed separately via Ansible because "
        "they touch customer VMs, not Azure resources.",
    )

    doc.add_heading("5.1  Workflow", level=2)
    add_numbered(doc, [
        "Clone the repository and change into the stack module directory.",
        "Copy terraform.tfvars.example to terraform.tfvars and fill in "
        "subscription_id and admin_password. Toggle deploy_vpb if you only "
        "want CLMS.",
        "Run terraform init to download the azurerm provider.",
        "Run terraform plan to review the resources Terraform will create. "
        "Confirm against your subscription's policies.",
        "Run terraform apply. About 8 to 15 minutes later, CLMS and vPB are "
        "ready. The outputs print the CLMS UI URL, the vPB management IP and "
        "the default credentials.",
        "Open the CLMS UI, create your project, copy the project key, and run "
        "quickstart.sh (or deploy-stack.sh --no-vpb) to deploy sensors.",
    ])

    add_code_block(
        doc,
        "git clone "
        "https://github.com/Keysight-Tech/cloudlens-ansible-azure.git\n"
        "cd cloudlens-ansible-azure/deploy/terraform/stack\n"
        "cp terraform.tfvars.example terraform.tfvars\n"
        "$EDITOR terraform.tfvars   # set subscription_id, admin_password\n"
        "terraform init\n"
        "terraform plan\n"
        "terraform apply",
    )

    doc.add_heading("5.2  What gets created", level=2)
    add_styled_table(
        doc,
        headers=["Resource", "Count", "Notes"],
        rows=[
            ["Resource group", "1", "cloudlens-rg (override via tfvars)"],
            ["Virtual network", "1", "cloudlens-stack-vnet 10.50.0.0/16"],
            ["Subnets", "4", "clms-subnet, vpb-mgmt, vpb-ingress, vpb-egress"],
            ["Public IPs", "2", "One for CLMS, one for vPB management"],
            ["NSGs", "2", "Pre-configured rules per VM role"],
            ["NICs", "4", "1 for CLMS, 3 for vPB"],
            ["VMs", "2", "CLMS D4s_v5, vPB D8s_v3"],
        ],
        col_widths=[1.8, 0.7, 3.7],
    )
    add_callout(
        doc, "Existing VNet?",
        "Set shared_vnet = false in terraform.tfvars. Each child module then "
        "creates its own VNet (CLMS gets a 1-subnet VNet, vPB gets a 3-subnet "
        "VNet). For full bring-your-own-VNet wiring, edit main.tf and pass "
        "existing_vnet_name through to the child modules.",
    )
    doc.add_page_break()


def build_path_c(doc: Document) -> None:
    doc.add_heading("6. Path C: Azure Portal", level=1)
    add_paragraph(
        doc,
        "If you prefer to deploy from the Portal, the same Marketplace images "
        "are also exposed via the two click-through buttons on the landing page "
        "at keysight-tech.github.io/cloudlens-ansible-azure. The Portal opens a "
        "Custom deployment blade that pulls a small wrapper ARM template "
        "directly from this repository.",
    )
    add_paragraph(
        doc,
        "Click the Deploy CLMS button first. The blade prompts for resource "
        "group, region, VM name, admin username and admin password. The "
        "deployment takes about 8 to 10 minutes; the VM then needs a further "
        "15 minutes to finish initialising before the UI is reachable.",
    )
    add_paragraph(
        doc,
        "Once CLMS is up, click the Deploy vPB button. Same blade pattern, "
        "different image. Use the same resource group as CLMS so both VMs end "
        "up in one tidy place that you can delete cleanly later.",
    )
    add_paragraph(
        doc,
        "Open the CLMS UI, create a project, copy the project key, and run "
        "quickstart.sh from a terminal (or Cloud Shell) to deploy sensors. The "
        "Portal does not yet have a click-through for sensors because they "
        "touch every tagged VM in the subscription; that is intentionally an "
        "explicit step.",
    )
    add_callout(
        doc, "Marketplace listing names",
        "In the Portal search, CLMS appears as 'CloudLens Manager (Preview)' "
        "and vPB appears as 'CloudLens Virtual Packet Broker'. The Deploy "
        "buttons on the site point at the correct offer IDs automatically, so "
        "you do not need to find them by hand.",
    )
    add_callout(
        doc, "No screenshots, by design",
        "Azure Portal UI shifts every few months. Rather than ship stale "
        "screenshots, this guide describes the click path in words. The "
        "actual buttons and field names match the Portal at June 2026.",
    )
    doc.add_page_break()


def build_verification(doc: Document) -> None:
    doc.add_heading("7. Verification Checklist", level=1)
    add_paragraph(
        doc,
        "Walk this list after every deploy. If any item fails, jump to section "
        "8 for the matching troubleshooting entry.",
    )
    add_checkbox_list(doc, [
        "CLMS UI reachable on https://<clms-public-ip>/. The page loads with "
        "the Keysight CloudLens login screen, not a connection error.",
        "Logged in to CLMS with the default credentials (admin / "
        "Cl0udLens@dm!n). Default password reset prompt appears; set a strong "
        "password before going further.",
        "Project created in CLMS under Settings > Projects, and the API key "
        "has been copied into a safe place.",
        "vPB SSH reachable on port 22 of the vPB management public IP after "
        "10 to 15 minutes. Login as admin with the admin password set during "
        "deploy.",
        "vPB CLI prompts confirm 'CloudLens vPB Console' on first SSH; reach "
        "config mode with 'enable' then 'configure'.",
        "Sensors visible in CLMS under Sensors page within 8 minutes of running "
        "quickstart.sh. Each sensor shows its hostname and Connected status.",
        "Test traffic from a sensored VM appears at the destination tool or "
        "probe (CLMS Investigate page or your downstream collector).",
    ])
    add_callout(
        doc, "Printable",
        "This page is intentionally a single checklist so it can be printed "
        "and ticked off during a customer handover or installation walkthrough.",
    )
    doc.add_page_break()


def build_troubleshooting(doc: Document) -> None:
    doc.add_heading("8. Troubleshooting Reference", level=1)
    add_paragraph(
        doc,
        "Organised by failure mode. The full reference (every error message "
        "we have seen in the field, plus the fix) lives at "
        "docs/TROUBLESHOOTING.md in the repository.",
    )

    doc.add_heading("Marketplace terms", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["MarketplacePurchaseEligibilityFailed during deploy",
             "Marketplace terms not accepted for the publisher / offer",
             "az vm image terms accept --publisher keysight-technologies-cloudlens "
             "--offer <offer> --plan <plan>"],
            ["Deploy succeeds but VM never starts",
             "Plan name typo (case-sensitive)",
             "Use exact plan ids: clms-6-13-0_76 and cloudlens-virtual-packet-broker-3-15-0_1"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    doc.add_heading("Quota", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["QuotaExceeded on DSv5 family",
             "Subscription default quota is too low",
             "Request quota increase in the Portal: Subscriptions > Usage + quotas; or pick a region with headroom"],
            ["Quota probe warns but deploy continues",
             "Quota is tight but enough for one VM",
             "Safe to ignore; the script does not block on quota warnings"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    doc.add_heading("Network", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["CLMS UI not reachable after 20 minutes",
             "NSG blocks port 443 from your client IP",
             "az network nsg rule create to allow your IP, or use Bastion"],
            ["vPB SSH times out",
             "NSG blocks port 22 from your client IP",
             "Add an NSG rule for your client IP on port 22"],
            ["Sensor cannot reach CLMS",
             "Egress blocked from VM subnet to CLMS public IP on 443",
             "Open egress on the workload subnet; or peer the VNets"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    doc.add_heading("Timing", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["CLMS UI returns 502 / 503 just after deploy",
             "CLMS is still initialising (services start in sequence)",
             "Wait the full 15 minutes; the script's poll loop handles this"],
            ["vPB SSH connection refused immediately after deploy",
             "vPB OS is still booting and applying first-run config",
             "Wait 10 minutes; if still refused, restart the VM from the Portal"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )
    doc.add_page_break()


def build_cleanup(doc: Document) -> None:
    doc.add_heading("9. Cleanup / Decommission", level=1)
    add_paragraph(
        doc,
        "Two ways to undo a stack deploy. Pick the one that matches how you "
        "deployed it.",
    )

    doc.add_heading("Terraform deploy", level=2)
    add_paragraph(
        doc,
        "If you used Path B (Terraform), use terraform destroy. This removes "
        "exactly what Terraform created and leaves anything else in the "
        "resource group untouched.",
    )
    add_code_block(
        doc,
        "cd deploy/terraform/stack\n"
        "terraform destroy",
    )

    doc.add_heading("Bash or Portal deploy", level=2)
    add_paragraph(
        doc,
        "If you used Path A or Path C, the simplest cleanup is to delete the "
        "whole resource group. This is irreversible; double-check the RG name "
        "before pressing enter.",
    )
    add_code_block(
        doc,
        "az group delete --name cloudlens-rg --yes --no-wait",
    )
    add_paragraph(
        doc,
        "The --no-wait flag returns immediately. The actual delete happens in "
        "the background and finishes in 1 to 3 minutes.",
    )

    doc.add_heading("Sensors on workload VMs", level=2)
    add_paragraph(
        doc,
        "Sensors are not destroyed by terraform destroy or az group delete, "
        "because they live on the customer's workload VMs (which are usually "
        "in a different resource group). To remove them:",
    )
    add_bullets(doc, [
        "Linux: docker stop cloudlens-agent && docker rm cloudlens-agent",
        "Windows: Stop-Service CloudLensAgent; sc.exe delete CloudLensAgent",
        "Or: untag the VMs (remove cloudlens=yes) and rerun the playbook with "
        "--tags remove",
    ])
    add_callout(
        doc, "Marketplace billing",
        "Deleting the resource group also stops the Keysight Marketplace "
        "subscription billing for those VMs at the next hourly tick. Compute "
        "billing stops immediately. Allow up to 24 hours for the Marketplace "
        "billing line to disappear from your invoice.",
    )
    doc.add_page_break()


def build_appendix(doc: Document) -> None:
    doc.add_heading("10. Appendix: File Paths and Quick Links", level=1)

    doc.add_heading("Repository file paths", level=2)
    add_styled_table(
        doc,
        headers=["File", "Purpose"],
        rows=[
            ["deploy/deploy-stack.sh", "Path A entry point (Bash one-liner)"],
            ["deploy/terraform/stack/", "Path B entry point (Terraform module)"],
            ["deploy/terraform/stack/terraform.tfvars.example", "Path B sample input"],
            ["deploy/clms-marketplace.json", "Wrapper ARM template for CLMS (used by Paths A and C)"],
            ["deploy/vpb-marketplace.json", "Wrapper ARM template for vPB (used by Paths A and C)"],
            ["docs/CloudLens_Stack_Deployment_Runbook.pdf", "This document (PDF)"],
            ["docs/CloudLens_Stack_Deployment_Runbook.docx", "This document (Word)"],
            ["quickstart.sh", "Sensor deployment chain (called by deploy-stack.sh)"],
        ],
        col_widths=[3.0, 4.0],
    )

    doc.add_heading("Command reference", level=2)
    add_styled_table(
        doc,
        headers=["Action", "Command"],
        rows=[
            ["Run full stack (recommended)",
             "curl -sSL https://raw.githubusercontent.com/Keysight-Tech/"
             "cloudlens-ansible-azure/main/deploy/deploy-stack.sh | bash"],
            ["Run full stack (dry-run)",
             "bash deploy/deploy-stack.sh --dry-run"],
            ["Run CLMS only (skip vPB)",
             "bash deploy/deploy-stack.sh --no-vpb"],
            ["Terraform apply",
             "cd deploy/terraform/stack && terraform apply"],
            ["Terraform destroy",
             "cd deploy/terraform/stack && terraform destroy"],
            ["Delete everything (RG)",
             "az group delete -n cloudlens-rg --yes --no-wait"],
        ],
        col_widths=[2.4, 4.6],
    )

    doc.add_heading("Useful links", level=2)
    add_styled_table(
        doc,
        headers=["Resource", "Location"],
        rows=[
            ["GitHub repository",
             "https://github.com/Keysight-Tech/cloudlens-ansible-azure"],
            ["Landing page",
             "https://keysight-tech.github.io/cloudlens-ansible-azure/"],
            ["Sensor runbook (companion document)",
             "docs/CloudLens_Ansible_Azure_Customer_Runbook.pdf"],
            ["Deployment guide (technical)",
             "docs/DEPLOYMENT_GUIDE.md"],
            ["Troubleshooting reference",
             "docs/TROUBLESHOOTING.md"],
            ["Issues / support",
             "https://github.com/Keysight-Tech/cloudlens-ansible-azure/issues"],
            ["Related: vPB + Azure GWLB",
             "https://github.com/Keysight-Tech/cloudlens-vpb-azure-gwlb"],
        ],
        col_widths=[2.6, 4.4],
    )
    add_paragraph(
        doc,
        "Support email subject (for the Keysight account team): "
        "'CloudLens Stack: <customer name>, <brief issue>'.",
        italic=True, size=10, color=TEXT_MUTED,
    )


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
def main() -> int:
    print(f"[info] writing {OUTPUT_DOCX}")

    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)

    build_cover_page(doc)
    build_toc(doc)
    build_executive_summary(doc)
    build_paths_overview(doc)
    build_prerequisites(doc)
    build_path_a(doc)
    build_path_b(doc)
    build_path_c(doc)
    build_verification(doc)
    build_troubleshooting(doc)
    build_cleanup(doc)
    build_appendix(doc)

    doc.save(OUTPUT_DOCX)
    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"[ok]  wrote {OUTPUT_DOCX.name} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
