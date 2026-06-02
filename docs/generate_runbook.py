#!/usr/bin/env python3
"""
CloudLens Ansible — Azure Customer Runbook generator.

Builds CloudLens_Ansible_Azure_Customer_Runbook.docx, mirroring the polished
README in an executive/printable format. Matches the visual style of the
existing GWLB customer runbook:
  - Keysight blue (#0078D4) section headers
  - Azure-blue header rows on tables
  - Courier New code blocks
  - Calibri body, professional sans-serif
  - Cover page + table of contents + page numbers (header/footer)

Embeds the four SVG assets in docs/assets/ by rasterising them to PNG via
cairosvg. If cairosvg is unavailable, falls back to a text caption so the
document still builds end-to-end.

Run:
    python3 docs/generate_runbook.py
Output:
    docs/CloudLens_Ansible_Azure_Customer_Runbook.docx
"""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# -----------------------------------------------------------------------------
# Brand palette
# -----------------------------------------------------------------------------
AZURE_BLUE = RGBColor(0x00, 0x78, 0xD4)
AZURE_DARK = RGBColor(0x00, 0x5A, 0x9E)
KEYSIGHT_GOLD = RGBColor(0xD4, 0xAF, 0x37)
KEYSIGHT_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
SUCCESS_GREEN = RGBColor(0x22, 0xC5, 0x5E)
TEXT_DARK = RGBColor(0x1B, 0x2A, 0x4A)
TEXT_MUTED = RGBColor(0x55, 0x65, 0x75)
CALLOUT_BG = "DDEBF7"  # light azure for callouts (hex string for shading)
CODE_BG = "F2F2F2"
TABLE_HEADER_BG = "0078D4"  # Azure blue
TABLE_ALT_ROW_BG = "F4F8FB"

# -----------------------------------------------------------------------------
# Paths
# -----------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
ASSETS_DIR = SCRIPT_DIR / "assets"
OUTPUT_DOCX = SCRIPT_DIR / "CloudLens_Ansible_Azure_Customer_Runbook.docx"

# -----------------------------------------------------------------------------
# SVG -> PNG conversion (graceful fallback if cairosvg missing)
# -----------------------------------------------------------------------------
try:
    import cairosvg  # type: ignore

    HAS_CAIROSVG = True
except Exception:  # pragma: no cover
    HAS_CAIROSVG = False


def svg_to_png(svg_path: Path, png_path: Path, width: int = 1400) -> bool:
    """Rasterise an SVG to PNG. Returns True on success."""
    if not HAS_CAIROSVG:
        return False
    if not svg_path.exists():
        return False
    try:
        cairosvg.svg2png(
            url=str(svg_path),
            write_to=str(png_path),
            output_width=width,
        )
        return True
    except Exception as exc:  # pragma: no cover
        print(f"[warn] cairosvg failed on {svg_path.name}: {exc}", file=sys.stderr)
        return False


# -----------------------------------------------------------------------------
# Low-level XML helpers
# -----------------------------------------------------------------------------
def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_MUTED


def _add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:t")
    fld_char3.text = (
        "Right-click and choose 'Update Field' to refresh the table of contents."
    )
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    run._r.append(fld_char4)


# -----------------------------------------------------------------------------
# Style configuration
# -----------------------------------------------------------------------------
def configure_styles(doc: Document) -> None:
    styles = doc.styles

    # Body / Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    # Heading 1 — Azure Blue
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = AZURE_BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — Azure Dark
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = AZURE_DARK
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3 — Keysight Navy
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = KEYSIGHT_NAVY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True


# -----------------------------------------------------------------------------
# Reusable element builders
# -----------------------------------------------------------------------------
def add_paragraph(doc: Document, text: str, *, bold: bool = False,
                  italic: bool = False, color: RGBColor | None = None,
                  size: int = 11, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_code_block(doc: Document, code: str) -> None:
    """Monospace Courier New block with light shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)

    # Shade paragraph background
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_BG)
    p_pr.append(shd)

    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = KEYSIGHT_NAVY


def add_callout(doc: Document, label: str, body: str) -> None:
    """Single-cell shaded callout box (Expected outcome / Note etc.)."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CALLOUT_BG)
    _set_cell_borders(cell, color=AZURE_BLUE_HEX, size="6")

    # Label
    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(2)
    run = label_p.add_run(label)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = AZURE_DARK

    # Body
    body_p = cell.add_paragraph()
    run2 = body_p.add_run(body)
    run2.font.name = "Calibri"
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = TEXT_DARK

    doc.add_paragraph()  # spacer


AZURE_BLUE_HEX = "0078D4"


def add_checkbox_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        box = p.add_run("☐  ")  # ballot box ☐
        box.font.name = "Segoe UI Symbol"
        box.font.size = Pt(13)
        box.font.color.rgb = AZURE_DARK
        text = p.add_run(item)
        text.font.name = "Calibri"
        text.font.size = Pt(11)
        text.font.color.rgb = TEXT_DARK


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                     col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade_cell(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            cell = row_cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                _shade_cell(cell, TABLE_ALT_ROW_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_DARK

    # Column widths (optional)
    if col_widths:
        for row in table.rows:
            for c_idx, width_in in enumerate(col_widths):
                row.cells[c_idx].width = Inches(width_in)

    doc.add_paragraph()


def add_image_or_placeholder(doc: Document, svg_name: str, caption: str,
                             width_inches: float = 6.5,
                             tmpdir: Path | None = None) -> None:
    svg_path = ASSETS_DIR / svg_name
    if HAS_CAIROSVG and tmpdir is not None and svg_path.exists():
        png_path = tmpdir / (svg_path.stem + ".png")
        if svg_to_png(svg_path, png_path, width=1600):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(png_path), width=Inches(width_inches))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(9.5)
            cap_run.font.color.rgb = TEXT_MUTED
            return
    # Fallback
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Diagram: {svg_name} — see docs/assets/{svg_name} ]")
    run.italic = True
    run.font.color.rgb = TEXT_MUTED
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9.5)
    cap_run.font.color.rgb = TEXT_MUTED


# -----------------------------------------------------------------------------
# Header / Footer
# -----------------------------------------------------------------------------
def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.different_first_page_header_footer = True

    # Header on inner pages
    header = section.header
    h_p = header.paragraphs[0]
    h_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = h_p.add_run("CloudLens Ansible — Azure Customer Runbook")
    h_run.font.name = "Calibri"
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = AZURE_DARK
    h_run.italic = True

    # Footer
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = f_p.add_run("Keysight Technologies   |   ")
    left.font.name = "Calibri"
    left.font.size = Pt(9)
    left.font.color.rgb = TEXT_MUTED
    _add_page_number(f_p)


# -----------------------------------------------------------------------------
# Cover page
# -----------------------------------------------------------------------------
def build_cover_page(doc: Document) -> None:
    # Top spacer
    for _ in range(3):
        doc.add_paragraph()

    # Logo placeholder
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run("KEYSIGHT")
    logo_run.font.name = "Calibri"
    logo_run.font.size = Pt(20)
    logo_run.bold = True
    logo_run.font.color.rgb = KEYSIGHT_GOLD

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline_p.add_run("TECHNOLOGIES")
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = KEYSIGHT_NAVY
    tag_run.bold = True

    for _ in range(4):
        doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CloudLens Ansible")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(36)
    title_run.bold = True
    title_run.font.color.rgb = AZURE_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Azure Customer Runbook")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(26)
    sub_run.font.color.rgb = AZURE_DARK

    doc.add_paragraph()
    doc.add_paragraph()

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag_p.add_run(
        "Automated sensor deployment for Linux and Windows VMs at scale"
    )
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(14)
    tag_run.italic = True
    tag_run.font.color.rgb = TEXT_MUTED

    for _ in range(8):
        doc.add_paragraph()

    # Version + date box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    for row in meta_table.rows:
        for cell in row.cells:
            cell.width = Inches(2.0)
    labels = [("Version", "v1.0.0"), ("Date", "June 2026")]
    for r_idx, (label, value) in enumerate(labels):
        l_cell = meta_table.rows[r_idx].cells[0]
        v_cell = meta_table.rows[r_idx].cells[1]
        _shade_cell(l_cell, "0078D4")
        _shade_cell(v_cell, "F4F8FB")
        lp = l_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        lr.font.name = "Calibri"
        lr.font.size = Pt(11)
        vp = v_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(value)
        vr.bold = True
        vr.font.color.rgb = KEYSIGHT_NAVY
        vr.font.name = "Calibri"
        vr.font.size = Pt(11)

    for _ in range(2):
        doc.add_paragraph()

    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot_p.add_run("Keysight Technologies   |   Network Visibility Solutions")
    foot_run.font.name = "Calibri"
    foot_run.font.size = Pt(10)
    foot_run.italic = True
    foot_run.font.color.rgb = TEXT_MUTED

    # Page break to next section
    doc.add_page_break()


# -----------------------------------------------------------------------------
# Section builders
# -----------------------------------------------------------------------------
def build_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    tip = doc.add_paragraph()
    tip_run = tip.add_run(
        "In Microsoft Word: right-click the entry below and choose "
        "'Update Field' to populate the full table of contents."
    )
    tip_run.italic = True
    tip_run.font.size = Pt(9.5)
    tip_run.font.color.rgb = TEXT_MUTED

    toc_p = doc.add_paragraph()
    _add_toc(toc_p)
    doc.add_page_break()


def build_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    add_paragraph(
        doc,
        "CloudLens Ansible for Azure is a fully automated kit that deploys "
        "Keysight CloudLens sensors to every tagged virtual machine in an Azure "
        "subscription — Linux and Windows, from a single VM to 5,000+ — without "
        "any manual per-host steps.",
    )
    add_paragraph(
        doc,
        "It is built for customer DevOps and network teams who need network "
        "visibility instrumented quickly and reliably, and for Keysight Sales "
        "Engineers who need a repeatable, demoable proof-of-value asset.",
    )

    doc.add_heading("Why it matters", level=2)
    bullets = [
        "Eliminates the per-VM SSH / RDP grind. Tag a VM with cloudlens=yes and the kit handles the rest.",
        "Three entry points — Azure Portal one-click, Cloud Shell curl, or Docker — so customers run it wherever they already work.",
        "Production-tested against real Azure. End-to-end deployment of 3 mixed-OS sensors verified in 8 minutes.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(b)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK

    doc.add_heading("Proof points", level=2)
    add_styled_table(
        doc,
        headers=["Scenario", "Result", "Time"],
        rows=[
            ["Ubuntu 22.04 (private IP via jumpbox)", "Sensor running", "4 min"],
            ["Ubuntu 22.04 (private IP via jumpbox)", "Sensor running", "4 min"],
            ["Windows Server 2022 (WinRM direct)", "Sensor running", "6 min"],
            ["CLMS 6.14.141 registration", "All 3 sensors registered", "<1 min"],
            ["End-to-end", "3/3 success", "8 min"],
        ],
        col_widths=[3.8, 2.0, 0.9],
    )
    add_paragraph(
        doc,
        "Date verified: 2026-06-02. Subscription: CloudLensPublic (eastus2).",
        italic=True, size=10, color=TEXT_MUTED,
    )
    doc.add_page_break()


def build_solution_overview(doc: Document, tmpdir: Path) -> None:
    doc.add_heading("2. Solution Overview", level=1)
    add_image_or_placeholder(
        doc, "architecture-diagram.svg",
        caption="Figure 1 — End-to-end architecture: control point → Azure inventory → OS lanes → CLMS",
        width_inches=6.5, tmpdir=tmpdir,
    )
    add_paragraph(
        doc,
        "A single Ansible control point authenticates to Azure (Service Principal, "
        "Managed Identity, or Cloud Shell session), queries the Azure Resource "
        "Manager API for every VM matching the customer's tag filter, and routes "
        "each host to the OS-specific playbook lane: Ubuntu, RHEL, or Windows.",
    )
    add_paragraph(
        doc,
        "Linux hosts run the CloudLens agent in a Docker container. Windows hosts "
        "run the CloudLens Windows sensor as a native service. Every sensor "
        "self-registers with CloudLens Manager (CLMS) on first start using the "
        "project key supplied in customer_input.yaml — no per-VM UI steps, no "
        "static inventory files to maintain.",
    )
    doc.add_page_break()


def build_choose_path(doc: Document, tmpdir: Path) -> None:
    doc.add_heading("3. Choosing Your Deployment Path", level=1)
    add_paragraph(
        doc,
        "All three paths run the same Ansible engine — same playbooks, same "
        "automation. Pick the entry point that matches how your team works.",
    )
    add_image_or_placeholder(
        doc, "decision-tree.svg",
        caption="Figure 2 — Decision tree: pick the entry point that matches your environment",
        width_inches=6.0, tmpdir=tmpdir,
    )

    add_styled_table(
        doc,
        headers=["Tier", "Best For", "Tools Needed", "Effort"],
        rows=[
            ["Tier 1: One-Click Portal",
             "Customers who live in the Azure Portal and want zero local setup",
             "Web browser only",
             "Lowest"],
            ["Tier 2: Cloud Shell",
             "Customers already authenticated to Azure in the browser",
             "Azure Cloud Shell (pre-authenticated)",
             "Low"],
            ["Tier 3: Docker",
             "Repeatable runs from a laptop or CI/CD pipeline",
             "Docker, Service Principal, customer_input.yaml",
             "Medium"],
        ],
        col_widths=[1.7, 2.5, 1.7, 0.9],
    )
    doc.add_page_break()


def build_prerequisites(doc: Document) -> None:
    doc.add_heading("4. Prerequisites Checklist (Printable)", level=1)
    add_paragraph(
        doc,
        "Confirm every item below before kicking off a deployment. Tick each "
        "box as you go — the kit will not magically fix a missing prerequisite.",
    )
    add_checkbox_list(doc, [
        "Azure subscription with at least Reader + Virtual Machine Contributor roles",
        "Service Principal created (Tier 3 only — see scripts/setup_azure_sp.sh)",
        "CLMS Manager deployed and reachable from the target VM subnets",
        "Project Key obtained from the CLMS UI (Projects → API Keys)",
        "Target VMs tagged with cloudlens=yes, os=ubuntu|rhel|windows, env=prod|dev|qa",
        "SSH key (Linux) and Windows admin password ready in your control environment",
    ])
    add_callout(
        doc,
        "Tip",
        "Tags are how the dynamic inventory discovers hosts. If a VM is not "
        "tagged, the kit cannot see it — and that is the single most common "
        "support ticket. Use the bulk-tag script in Appendix B to tag an entire "
        "resource group in one shot.",
    )
    doc.add_page_break()


def build_deployment(doc: Document) -> None:
    doc.add_heading("5. Deployment — Step by Step", level=1)

    # ---- Tier 1 ----
    doc.add_heading("5.1  Tier 1: One-Click from the Azure Portal", level=2)
    add_paragraph(
        doc,
        "Best when you want zero local tools. The Portal provisions a short-lived "
        "Ubuntu runner VM in your subscription, the runner uses Managed Identity "
        "to authenticate, deploys sensors to every tagged VM, then self-destructs.",
    )
    t1_steps = [
        "Open the README in GitHub and click the 'Deploy to Azure' button.",
        "Sign in to the Azure Portal when prompted; the ARM template opens in a 'Custom deployment' blade.",
        "Fill in the four parameters: CLMS IP/FQDN, Project Key, Tag Filter (default cloudlens=yes), and Resource Group for the runner.",
        "Review + Create. Azure provisions the runner VM (~2 minutes).",
        "Watch the runner's boot diagnostics log — every tagged VM is sensored in parallel.",
        "Runner self-destructs after 1 hour. Confirm sensors in CLMS → Sensors page.",
    ]
    for idx, step in enumerate(t1_steps, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    add_paragraph(doc, "[Screenshot placeholder — Azure Portal Deploy blade]",
                  italic=True, size=10, color=TEXT_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc, "Expected outcome",
        "Tagged VMs appear in CLMS → Sensors within 8 minutes. Runner VM is "
        "destroyed automatically. No artefacts left on the customer's laptop.",
    )

    # ---- Tier 2 ----
    doc.add_heading("5.2  Tier 2: Azure Cloud Shell", level=2)
    add_paragraph(
        doc,
        "Best when the customer is already logged into Azure in a browser tab. "
        "Cloud Shell is pre-authenticated, so there is no Service Principal to "
        "create and no local install.",
    )
    t2_steps = [
        "Open Azure Cloud Shell (shell.azure.com) and choose Bash.",
        "Run the quickstart bootstrap:",
        "Answer the wizard prompts: CLMS IP, Project Key, tag filter, connection mode.",
        "Confirm sensors in CLMS → Sensors page (~8 min for a small environment).",
    ]
    for idx, step in enumerate(t2_steps, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    add_code_block(
        doc,
        "curl -sSL https://raw.githubusercontent.com/Keysight-Tech/"
        "cloudlens-ansible-azure/main/quickstart.sh | bash",
    )
    add_callout(
        doc, "Expected outcome",
        "Wizard discovers tagged VMs, auto-tunes Ansible forks, deploys sensors, "
        "and prints a 3/3 success summary. All state lives in your Cloud Shell "
        "home directory — nothing installed locally.",
    )

    # ---- Tier 3 ----
    doc.add_heading("5.3  Tier 3: Docker (laptop or CI/CD)", level=2)
    add_paragraph(
        doc,
        "Best for repeatable runs from a developer laptop, a CI pipeline, or "
        "any container host. The image is pinned and hermetic — same result on "
        "macOS, Windows, Linux, GitHub Actions, GitLab CI, and Jenkins.",
    )
    t3_steps = [
        "Create a Service Principal with `bash scripts/setup_azure_sp.sh` (one-time).",
        "Copy customer_input.yaml.example to customer_input.yaml and fill in your CLMS IP, project key, and tag filter.",
        "Export AZURE_SUBSCRIPTION_ID, AZURE_TENANT, AZURE_CLIENT_ID, AZURE_SECRET, ANSIBLE_WINRM_PASSWORD as env vars.",
        "Run the container:",
    ]
    for step in t3_steps:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    add_code_block(
        doc,
        "docker run --rm -it \\\n"
        "  -v $(pwd)/customer_input.yaml:/work/customer_input.yaml \\\n"
        "  -v $HOME/.ssh:/root/.ssh:ro \\\n"
        "  -e AZURE_SUBSCRIPTION_ID -e AZURE_TENANT \\\n"
        "  -e AZURE_CLIENT_ID -e AZURE_SECRET \\\n"
        "  -e ANSIBLE_WINRM_PASSWORD \\\n"
        "  ghcr.io/keysight-tech/cloudlens-ansible-azure:latest",
    )
    add_callout(
        doc, "Expected outcome",
        "Container exits 0 with a 'sensors deployed: N/N' summary. Pin the "
        "image tag in CI to get hermetic, reproducible runs.",
    )
    doc.add_page_break()


def build_scenarios(doc: Document, tmpdir: Path) -> None:
    doc.add_heading("6. Supported VM Scenarios", level=1)
    add_image_or_placeholder(
        doc, "scenario-matrix.svg",
        caption="Figure 3 — VM compatibility matrix (OS × topology × auth method)",
        width_inches=6.5, tmpdir=tmpdir,
    )
    add_styled_table(
        doc,
        headers=["OS / Topology", "Public IP direct", "Private + Jumpbox",
                 "Azure Bastion", "Cloud Shell"],
        rows=[
            ["Ubuntu 20.04 / 22.04 / 24.04", "Supported", "Supported", "Supported", "Supported"],
            ["RHEL 7 / 8 / 9", "Supported", "Supported", "Supported", "Supported"],
            ["CentOS / Rocky / AlmaLinux", "Supported", "Supported", "Supported", "Supported"],
            ["Windows Server 2019 / 2022", "Supported", "Planned", "Planned", "Supported"],
        ],
        col_widths=[2.4, 1.2, 1.3, 1.0, 1.1],
    )
    doc.add_page_break()


def build_verification(doc: Document) -> None:
    doc.add_heading("7. Verification Checklist (Printable)", level=1)
    add_paragraph(
        doc,
        "After every deployment, walk this list. If any item fails, jump to "
        "section 9 (Troubleshooting Reference).",
    )
    add_checkbox_list(doc, [
        "All sensors visible in CLMS → Sensors page",
        "Filter by custom_tags (e.g. Customer=Acme) returns the expected VMs",
        "Test traffic from a sensored VM shows up in the destination tool/probe",
        "Container logs show no errors (Linux): `docker logs cloudlens-agent`",
        "CloudLens service running (Windows): `Get-Service CloudLensAgent`",
    ])
    doc.add_page_break()


def build_scaling(doc: Document) -> None:
    doc.add_heading("8. Scaling Guide", level=1)
    add_styled_table(
        doc,
        headers=["VM Count", "Auto Forks", "Sharded?", "Approx Time"],
        rows=[
            ["1–50", "20", "No", "5–10 min"],
            ["50–500", "50", "No", "15–30 min"],
            ["500–2,000", "200", "No", "30–60 min"],
            ["2,000–10,000", "500 / shard", "Yes (auto)", "30–60 min"],
            ["10,000+", "1000 / shard", "AWX", "1–2 hr"],
        ],
        col_widths=[1.4, 1.4, 1.6, 1.6],
    )
    add_paragraph(
        doc,
        "The kit auto-tunes the Ansible fork count based on the number of VMs "
        "the dynamic inventory discovers. Above 2,000 VMs it also auto-shards: "
        "the inventory is split into batches of `deploy.shard_size` VMs and the "
        "playbook runs them sequentially. This keeps memory and Azure API "
        "throttling under control while still finishing in under an hour for "
        "tens of thousands of hosts. See docs/SCALING.md in the repository for "
        "AWX/Tower integration details.",
    )
    doc.add_page_break()


def build_troubleshooting(doc: Document) -> None:
    doc.add_heading("9. Troubleshooting Reference", level=1)
    add_paragraph(
        doc,
        "Organised by symptom area. The full reference lives at "
        "docs/TROUBLESHOOTING.md in the repository.",
    )

    doc.add_heading("Inventory", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["Inventory finds 0 VMs",
             "Tags missing on VMs",
             "az vm update --set tags.cloudlens=yes tags.os=<os> tags.env=prod"],
            ["Subset of VMs missing",
             "Tag value mismatch (case-sensitive)",
             "Use exact lowercase values: ubuntu, rhel, windows, prod"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )

    doc.add_heading("SSH (Linux)", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["Permission denied (publickey)",
             "Public key not present on target VM",
             "Bootstrap via az vm run-command invoke or use jumpbox mode"],
            ["Connection timeout",
             "NSG blocks port 22 from control point",
             "Switch connection.mode to jumpbox or bastion"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )

    doc.add_heading("WinRM (Windows)", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["WinRM timeout on 5985",
             "WinRM disabled on Windows VM",
             "Run playbooks/bootstrap_windows_winrm.yaml against the host"],
            ["401 Unauthorized",
             "ANSIBLE_WINRM_PASSWORD env var not exported",
             "export ANSIBLE_WINRM_PASSWORD='...' before running the kit"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )

    doc.add_heading("Container (Linux sensor)", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["apt_pkg.Error: Signed-By conflict",
             "Stale Docker apt source list",
             "The playbook auto-cleans on the next run; rerun the kit"],
            ["Container restarts in a loop",
             "Wrong CLMS IP or project key",
             "docker logs cloudlens-agent; fix manager_ip_or_fqdn / project_key"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )

    doc.add_heading("Registration with CLMS", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["Sensor not in CLMS UI",
             "Wrong project key",
             "Check CLMS → Projects → API Keys; update customer_input.yaml"],
            ["Sensor offline after restart",
             "CLMS IP unreachable from VM subnet",
             "Open egress from VM subnet to CLMS on TCP/443"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )
    doc.add_page_break()


def build_appendix_a(doc: Document) -> None:
    doc.add_heading("Appendix A — customer_input.yaml Schema", level=1)
    add_paragraph(
        doc,
        "Full annotated example. Copy customer_input.yaml.example from the "
        "repository to customer_input.yaml and fill in values. Never commit "
        "customer_input.yaml to git — credentials belong in env vars.",
    )
    add_code_block(doc, """# === Azure Environment ===
azure:
  subscription_id: "00000000-0000-0000-0000-000000000000"
  tenant_id:       "00000000-0000-0000-0000-000000000000"

  # Tag selector — VMs are matched by these tags
  tag_filters:
    cloudlens: "yes"
    env:       "prod"

  resource_groups: []   # empty = all RGs in the subscription
  locations: []         # empty = all regions

# === CloudLens Configuration ===
cloudlens:
  manager_ip_or_fqdn: "clms.customer.example.com"
  project_key:        "REPLACE_WITH_PROJECT_KEY"
  custom_tags:        "Env=Azure Region=eastus2 Customer=Acme"
  registry_type:      "insecure"   # or "secure"
  ssl_verify:         "no"
  auto_update:        "yes"
  local_ca_path:      "files/cloudlenscerts.crt"
  ca_cert_dir:        "/etc/ssl/certs"
  agent_container_name: "cloudlens-agent"
  log_max_size: "50m"
  log_max_file: "5"

# === Connection Mode ===
# direct_public | jumpbox | bastion | cloud_shell
connection:
  mode: "jumpbox"
  jumpbox_host:    "1.2.3.4"
  jumpbox_user:    "azureuser"
  jumpbox_ssh_key: "~/.ssh/id_rsa"
  bastion_resource_id: ""

# === Linux SSH ===
linux:
  ansible_user: "azureuser"
  ssh_key_file: "~/.ssh/id_rsa"

# === Windows WinRM ===
# Password MUST be set via env var: export ANSIBLE_WINRM_PASSWORD='...'
windows:
  ansible_user:                       "azureuser"
  ansible_connection:                 "winrm"
  ansible_winrm_transport:            "basic"
  ansible_winrm_server_cert_validation: "ignore"
  ansible_port:                       5985
  installer_path:     "files/cloudlens-win-sensor-6.13.0.359.exe"
  installer_filename: "cloudlens-win-sensor-6.13.0.359.exe"

# === Deployment Behavior ===
deploy:
  forks: 0                # 0 = auto-tune
  timeout_seconds: 60
  reinstall_if_unhealthy: true
  shard_size: 500         # for >2000 VMs
""")
    doc.add_page_break()


def build_appendix_b(doc: Document) -> None:
    doc.add_heading("Appendix B — Bulk Tag Script", level=1)
    add_paragraph(
        doc,
        "Apply the three required tags to every VM in a resource group, "
        "filtered by OS. Run these from Azure Cloud Shell or any shell with "
        "the Azure CLI logged in.",
    )

    doc.add_heading("Ubuntu VMs", level=3)
    add_code_block(doc, """for vm in $(az vm list -g <RG> --query \\
    "[?storageProfile.imageReference.offer=='0001-com-ubuntu-server-jammy'].name" \\
    -o tsv); do
  az vm update -g <RG> -n $vm \\
    --set tags.cloudlens=yes tags.os=ubuntu tags.env=prod
done""")

    doc.add_heading("RHEL VMs", level=3)
    add_code_block(doc, """for vm in $(az vm list -g <RG> --query \\
    "[?storageProfile.imageReference.publisher=='RedHat'].name" \\
    -o tsv); do
  az vm update -g <RG> -n $vm \\
    --set tags.cloudlens=yes tags.os=rhel tags.env=prod
done""")

    doc.add_heading("Windows VMs", level=3)
    add_code_block(doc, """for vm in $(az vm list -g <RG> --query \\
    "[?storageProfile.osDisk.osType=='Windows'].name" \\
    -o tsv); do
  az vm update -g <RG> -n $vm \\
    --set tags.cloudlens=yes tags.os=windows tags.env=prod
done""")
    doc.add_page_break()


def build_appendix_c(doc: Document) -> None:
    doc.add_heading("Appendix C — Quick Links", level=1)
    add_styled_table(
        doc,
        headers=["Resource", "Location"],
        rows=[
            ["GitHub repository", "https://github.com/Keysight-Tech/cloudlens-ansible-azure"],
            ["README (technical entry point)",
             "https://github.com/Keysight-Tech/cloudlens-ansible-azure/blob/main/README.md"],
            ["Deployment guide",
             "docs/DEPLOYMENT_GUIDE.md (in the repo)"],
            ["Scaling guide",
             "docs/SCALING.md (in the repo)"],
            ["Troubleshooting guide",
             "docs/TROUBLESHOOTING.md (in the repo)"],
            ["Customer email template",
             "docs/CUSTOMER_EMAIL.md (in the repo)"],
            ["GitHub Issues (bug reports / feature requests)",
             "https://github.com/Keysight-Tech/cloudlens-ansible-azure/issues"],
            ["Related: vPB + Azure GWLB",
             "https://github.com/Keysight-Tech/cloudlens-vpb-azure-gwlb"],
        ],
        col_widths=[2.6, 4.4],
    )
    add_paragraph(
        doc,
        "Support email subject (for the Keysight account team): "
        "'CloudLens Ansible Azure — <customer name> — <brief issue>'.",
        italic=True, size=10, color=TEXT_MUTED,
    )


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
def main() -> int:
    if not HAS_CAIROSVG:
        print("[warn] cairosvg unavailable — diagrams will be rendered as captions.")
    print(f"[info] writing {OUTPUT_DOCX}")

    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)

    build_cover_page(doc)

    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        build_toc(doc)
        build_executive_summary(doc)
        build_solution_overview(doc, tmpdir)
        build_choose_path(doc, tmpdir)
        build_prerequisites(doc)
        build_deployment(doc)
        build_scenarios(doc, tmpdir)
        build_verification(doc)
        build_scaling(doc)
        build_troubleshooting(doc)
        build_appendix_a(doc)
        build_appendix_b(doc)
        build_appendix_c(doc)

        doc.save(OUTPUT_DOCX)

    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"[ok]  wrote {OUTPUT_DOCX.name} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
